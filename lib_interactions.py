import nltk
import docx
import json
import copy
import spacy
import numpy
from nltk.corpus import stopwords
from nltk.parse.malt import MaltParser
import nltk.corpus as corp


UNSHIPHER = {
    "CC": "союз",
    "CD": "кардинальное число",
    "DT": "определитель",
    "EX": 'существование там',
    "FW": "иностранное слово",
    "IN": "предлог/подчинительный союз",
    "JJ": 'прилагательное',
    "VP": "глагольная группа",
    "JJR": 'прилагательное, сравнительная степень',
    "JJS": 'прилагательное, превосходная степень',
    "LS": "маркер списка  1)",
    "MD": "модальный глагол сосотавное сказуемое",
    "NN": 'существительное, единственное число',
    "NNS": 'существительное, множественное число',
    "PP": "предложная группа",
    "NNP": 'имя собственное, единственное число',
    "NNPS": 'имя собственное, множественное число',
    "PDT": 'предопределитель',
    "POS": "притяжательное окончание",
    "PRP": "личное местоимение, ",
    "PRP$": "притяжательное местоимение",
    "RB": "наречие",
    "RBR": "наречие, сравнительная степень",
    "RBS": "наречие, превосходная степень",
    "RP": "частица",
    "S": "Простое повествовательное предложение",
    "SBAR": "Предложение, введенное (возможно пустым) подчинительным союзом",
    "SBARQ": "Прямой вопрос, введенный вопросительным словом или вопросительной группой",
    "SINV": "Инвертированное повествовательное предложение, т.е. такое, в котором подлежащее следует за глаголом в прошедшем времени или модальным глаголом.",
    "SQ": "Инвертированный вопрос да/нет, или главное предложение вопроса, следующее за вопросительной группой в SBARQ",
    "SYM": "Символ",
    "VBD": "глагол, прошедшее время",
    "VBG": "глагол, герундий/презенс-партицип  берущий",
    "VBN": "глагол, прошедшее причастие  взятый",
    "VBP": "глагол, настоящее время, ед. число, не 3-е лицо",
    "VBZ": "глагол, настоящее время, 3-е лицо, ед. число",
    "WDT": "вопросительный определитель",
    "WP": "вопросительное местоимение",
    "WP$": "притяжательное вопросительное местоимение",
    "WRB": "вопросительное наречие",
    "TO": 'to',
    "UH": "междометие",
    "VB": "глагол, исходная форма",
}


signs = "!~@#$%^&*()_+<>?:.,;[]\\|'\"\'–«‘1234567890"


def load_file(path: str) -> list:
    opened_file = open(path, "rb")
    if '.doc' in path:
        current_file = docx.Document(opened_file)
        list_of_rows = [row.text for row in current_file.paragraphs]
        opened_file.close()
        raw_data = ""
        for row in list_of_rows:
            raw_data += row
    elif '.txt' in path:
        raw_data = ''
        for raw in opened_file:
            raw_data += raw
    else:
        raw_data = None
    return raw_data


def process_text(raw_data):
    stop_words = set(stopwords.words("english"))
    corpus_file = open('corpus.json', 'r')
    corpus = json.load(corpus_file)
    corpus_file.close()
    for sign in signs:
        stop_words.add(sign)
    raw_data = raw_data.replace('.', ' ')
    word_tokens = nltk.tokenize.word_tokenize(raw_data)
    tagged_word_tokens = nltk.pos_tag(word_tokens)
    final_table = []
    for token in tagged_word_tokens:
        if (token[0].lower() not in stop_words) and not token[0].isdigit():
            final_table.append(token)
    for token in range(len(final_table)):
        buffer = final_table[token][0]
        final_table[token] = (buffer, UNSHIPHER.get(final_table[token][1]))
    final_table = set(final_table)
    final_table = list(final_table)
    elemenation_list = []
    for entry in range(len(final_table)):
        if (final_table[entry][0] not in corpus.keys()) or not final_table[entry][1]:
            elemenation_list.append(final_table[entry])
        else:
            final_table[entry] = (final_table[entry][0], final_table[entry][1] +
                                  ' ,количество вхождений: ' + str(corpus.get(final_table[entry][0])))
    for entry in elemenation_list:
        final_table.remove(entry)
    return final_table


def save_file(text: list, name: str):
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "id"
    hdr_cells[1].text = "word"
    hdr_cells[2].text = "description"
    for id, word, description in text:
        if not description:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = str(id)
        row_cells[1].text = word
        row_cells[2].text = description

    document.add_page_break()

    document.save(name + ".docx")
